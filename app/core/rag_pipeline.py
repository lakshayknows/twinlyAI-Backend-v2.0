# app/core/rag_pipeline.py

import os
import json
import logging
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
import shutil

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langgraph.prebuilt import create_react_agent
from langchain_core.tools import tool
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, AIMessage
from app.core.config import settings

from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field
from typing import List

# --- GLOBAL MODEL CACHE ---
_EMBEDDINGS_MODEL = None

def get_embeddings_model():
    global _EMBEDDINGS_MODEL
    if _EMBEDDINGS_MODEL is None:
        _EMBEDDINGS_MODEL = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    return _EMBEDDINGS_MODEL

# --- Pydantic model for metadata extraction ---
class ResumeMetadata(BaseModel):
    candidate_name: str = Field(description="The full name of the candidate")
    summary: str = Field(description="A concise 2-sentence professional summary of the candidate")
    skills: List[str] = Field(description="A list of the top 10 most relevant technical skills or tools")
    experience_years: float = Field(description="Total estimated years of professional experience as a number (e.g., 3.5)")

class InterviewAssessment(BaseModel):
    topics_covered: List[str] = Field(description="High-level topics or skills discussed so far")
    red_flags: List[str] = Field(description="Any concerns or missing skills that surfaced during the chat")
    recruiter_intent: str = Field(description="What the recruiter seems to be looking for or prioritizing")

# --- JSON to TEXT CONVERSION (Helper Function) ---
def json_to_text(json_data: dict) -> str:
    text = ""
    for key, value in json_data.items():
        if isinstance(value, dict):
            text += "{}:\n".format(key.replace('_', ' ').title())
            for sub_key, sub_value in value.items():
                text += "  {}: {}\n".format(sub_key.replace('_', ' ').title(), sub_value)
        elif isinstance(value, list):
            text += "{}:\n".format(key.replace('_', ' ').title())
            for item in value:
                if isinstance(item, dict):
                    for item_key, item_value in item.items():
                        text += "  - {}: {}\n".format(item_key.replace('_', ' ').title(), item_value)
                else:
                    text += "- {}\n".format(item)
        else:
            text += "{}: {}\n".format(key.replace('_', ' ').title(), value)
    return text

# --- FILE PROCESSING (Helper Function) ---
def extract_text_from_file(file_path: Path) -> str:
    if file_path.suffix == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            return "".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_path.suffix == ".docx":
        doc = DocxDocument(file_path)
        return "\n".join(para.text for para in doc.paragraphs)
    elif file_path.suffix == ".txt":
        return file_path.read_text(encoding="utf-8")
    elif file_path.suffix == ".json":
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json_to_text(data)
    else:
        raise ValueError("Unsupported file type: {}".format(file_path.suffix))

class RAGPipeline:
    def __init__(self, bot_id: str, user_id: str, bot_name: str):
        self.bot_id = bot_id
        self.user_id = user_id
        self.bot_name = bot_name
        self.data_path = Path("data") / user_id / bot_id
        
        self.embeddings = get_embeddings_model()
        self.qdrant_client = QdrantClient(url=settings.QDRANT_URL, api_key=settings.QDRANT_API_KEY)
        self.collection_name = "bot_{}".format(self.bot_id)
        
        self.llm = ChatGroq(
            model_name="llama-3.3-70b-versatile", 
            temperature=0.7, 
            groq_api_key=settings.GROQ_API_KEY
        )
        
        self.vector_store = self._load_vector_store()
        self.agent_executor = None # Will instantiate dynamically based on metadata

    def _load_vector_store(self):
        if self.qdrant_client.collection_exists(self.collection_name):
            try:
                return QdrantVectorStore(
                    client=self.qdrant_client, 
                    collection_name=self.collection_name, 
                    embedding=self.embeddings
                )
            except Exception as e:
                print("Error loading Qdrant vector store")
                logging.exception("Error loading Qdrant vector store")
                return None
        return None

    def _create_agent(self, dynamic_metadata_text: str = ""):
        system_prompt = """You are \"{bot_name}\", a professional AI assistant representing a candidate.
You answer recruiter questions based on the candidate's resume.

Here is the candidate's professional profile:
---
{metadata}
---

Rules:
1. Only answer from the resume context. Say \"I don't have that in my resume\" if unsure.
2. Be professional and natural.
3. If the question is a greeting, greet back warmly.
4. Keep answers concise.""".format(bot_name=self.bot_name, metadata=dynamic_metadata_text)
        
        # --- DEFINE TOOLS ---
        @tool
        def search_resume(query: str) -> str:
            """Search the candidate's resume for specific details, work history, or keywords."""
            if not self.vector_store:
                return "Resume vector store not initialized."
            
            base_retriever = self.vector_store.as_retriever(search_kwargs={"k": 5})
            docs = base_retriever.invoke(query)
            return "\n\n".join([doc.page_content for doc in docs])
        
        @tool
        def calculate_experience(start_year: int, end_year: int) -> int:
            """Mathematically compute the number of years between two dates to confirm experience."""
            return end_year - start_year
        
        tools = [search_resume, calculate_experience]
        
        return create_react_agent(
            model=self.llm,
            tools=tools,
            prompt=system_prompt
        )

    def process_file(self, file_path: str):
        text_content = extract_text_from_file(Path(file_path))
        documents = [Document(page_content=text_content)]
        
        # --- RECURSIVE CHARACTER TEXT SPLITTER ---
        # Faster and avoids excessive API calls during embedding
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200, length_function=len)
        splits = text_splitter.split_documents(documents)

        # Build Qdrant store
        self.vector_store = QdrantVectorStore.from_documents(
            documents=splits, 
            embedding=self.embeddings,
            url=settings.QDRANT_URL,
            api_key=settings.QDRANT_API_KEY,
            collection_name=self.collection_name
        )
        
        # Agent will be constructed dynamically during stream response to include DB metadata
        return True
        
    async def extract_metadata(self, file_path: str) -> dict:
        """
        Uses the Maverick model to extract structured metadata (skills, exp, summary) from the resume.
        """
        text_content = extract_text_from_file(Path(file_path))
        
        # Truncate text to avoid token limits if resume is huge
        truncated_text = text_content[:12000] 

        parser = JsonOutputParser(pydantic_object=ResumeMetadata)

        prompt = ChatPromptTemplate.from_messages([
            ("system", "You are an expert technical recruiter and data analyst. Your task is to extract structured data from the following resume text. You must return ONLY a valid JSON object. Do not add any conversational text or markdown formatting around the JSON."),
            ("human", "Resume Text:\n{resume_text}\n\n{format_instructions}")
        ])
        
        extraction_llm = ChatGroq(
            model_name="llama-3.3-70b-versatile",
            temperature=0.0, 
            groq_api_key=settings.GROQ_API_KEY
        )

        chain = prompt | extraction_llm | parser

        try:
            metadata = await chain.ainvoke({
                "resume_text": truncated_text,
                "format_instructions": parser.get_format_instructions()
            })
            return metadata
        except Exception as e:
            logging.exception("Error extracting metadata")
            return {
                "candidate_name": self.bot_name,
                "summary": "Summary could not be extracted.",
                "skills": [],
                "experience_years": 0.0
            }

    async def analyze_interview(self, chat_history: list) -> dict:
        """Analyzes the current chat history to extract interview state."""
        if len(chat_history) < 2:
            return {} # Not enough context
        
        # Take the last 10 messages for context so we don't blow up token limits
        recent_history = chat_history[-10:]
        history_text = "\n".join(["{}: {}".format(getattr(msg, 'type', 'unknown'), msg.content) for msg in recent_history])
        
        parser = JsonOutputParser(pydantic_object=InterviewAssessment)
        prompt = ChatPromptTemplate.from_messages([
            ("system", "You are an expert HR recruiter assistant analyzing a chat between a recruiter and an AI representing a candidate. Extract the current state of the interview. Always return valid JSON.\n\n{format_instructions}"),
            ("human", "Chat History:\n{history_text}")
        ])
        
        extraction_llm = ChatGroq(
            model_name="llama-3.3-70b-versatile",
            temperature=0.0, 
            groq_api_key=settings.GROQ_API_KEY
        )
        chain = prompt | extraction_llm | parser
        try:
            result = await chain.ainvoke({
                "history_text": history_text,
                "format_instructions": parser.get_format_instructions()
            })
            return result
        except Exception as e:
            logging.exception("Error analyzing interview")
            return {}

    async def get_response_stream(self, user_message: str, chat_history: list = [], bot_metadata: dict = None):
        projects_text = ""
        links_text = ""
        meta_context = ""

        if bot_metadata:
            # Build Projects String
            projects = bot_metadata.get('projects', [])
            if isinstance(projects, list) and len(projects) > 0:
                projects_text = "\nFeatured Projects:\n"
                for p in projects:
                    proj_name = p.get('name', 'Unnamed')
                    proj_desc = p.get('description', '')
                    proj_link = p.get('link', '')
                    projects_text += "- {}: {} ({})\n".format(proj_name, proj_desc, proj_link)

            # Build Links String
            links_text += "LinkedIn: {}\n".format(bot_metadata.get('linkedin_url', 'N/A'))
            links_text += "GitHub: {}\n".format(bot_metadata.get('github_url', 'N/A'))
            links_text += "Twitter: {}\n".format(bot_metadata.get('twitter_url', 'N/A'))
            links_text += "Website: {}\n".format(bot_metadata.get('website_url', 'N/A'))

            meta_context = (
                "Name: {}\n"
                "Summary: {}\n"
                "Skills: {}\n"
                "Experience: {} years\n"
                "{}"
                "{}"
            ).format(
                bot_metadata.get('name', self.bot_name),
                bot_metadata.get('summary', 'Not available'),
                ', '.join(bot_metadata.get('skills') or []),
                bot_metadata.get('experience_years', 'Unknown'),
                links_text,
                projects_text
            )

        # Ensure agent is instantiated with latest metadata
        if self.vector_store:
            self.agent_executor = self._create_agent(dynamic_metadata_text=meta_context)

        messages = chat_history + [HumanMessage(content=user_message)]

        if not self.agent_executor:
            # Fallback: answer purely from metadata stored in MongoDB when Qdrant has no vectors
            if bot_metadata:
                fallback_prompt = ChatPromptTemplate.from_messages([
                    ("system", """You are \"{bot_name}\", a professional AI assistant representing a candidate.
Answer questions about this person based ONLY on the information below. Speak in third person.
If asked something not covered, politely say it's not available.

Candidate Profile:
{meta_context}"""),
                    MessagesPlaceholder(variable_name="chat_history"),
                    ("human", "{input}"),
                ])
                chain = fallback_prompt | self.llm
                async for chunk in chain.astream({"input": user_message, "chat_history": chat_history}):
                    if hasattr(chunk, "content"):
                        yield chunk.content
            else:
                yield "Error: The AI bot has not been properly initialized. Please upload a resume."
            return

        async for event in self.agent_executor.astream_events(
            {"messages": messages}, 
            version="v2"
        ):
            kind = event["event"]
            if kind == "on_chat_model_stream":
                content = event["data"]["chunk"].content
                if content:
                    yield content

# --- GLOBAL RECRUITER INDEX (SEMANTIC SEARCH) ---
class GlobalRecruiterIndex:
    """
    Manages a global Qdrant index that stores a summary profile for EVERY candidate
    to enable semantic search across the entire talent pool.
    """
    def __init__(self):
        self.collection_name = "global_recruiters_index"
        self.embeddings = get_embeddings_model()
        self.qdrant_client = QdrantClient(url=settings.QDRANT_URL, api_key=settings.QDRANT_API_KEY)

    def add_candidate_profile(self, bot_id: str, profile_text: str):
        """
        Adds or updates a candidate's profile in the global search index.
        """
        doc = Document(page_content=profile_text, metadata={"bot_id": bot_id})
        
        # Use from_documents which will create or update the collection
        QdrantVectorStore.from_documents(
            documents=[doc],
            embedding=self.embeddings,
            url=settings.QDRANT_URL,
            api_key=settings.QDRANT_API_KEY,
            collection_name=self.collection_name
        )
        return True

    def semantic_search(self, query: str, k: int = 10) -> List[str]:
        """
        Performs a semantic search and returns a list of matching bot_ids.

        Bypasses LangChain's QdrantVectorStore for the search query to avoid
        hybrid-vector name-mismatch issues. Instead:
          1. Embeds the query with HuggingFace directly
          2. Queries Qdrant using the unnamed dense vector (key='')
        """
        import logging as _log

        if not self.qdrant_client.collection_exists(self.collection_name):
            logging.info("[GlobalRecruiterIndex] Collection '%s' does not exist.", self.collection_name)
            return []

        try:
            # Step 1: Embed the query using HuggingFace Inference API
            logging.debug("[GlobalRecruiterIndex] Embedding query: '%s'", query[:80])
            query_vector = self.embeddings.embed_query(query)
            logging.debug("[GlobalRecruiterIndex] Query embedded successfully, dim=%d", len(query_vector))
        except Exception as e:
            logging.exception("[GlobalRecruiterIndex] Embedding failed")
            raise RuntimeError("Failed to embed search query.") from e

        try:
            # Step 2: Query Qdrant directly using the dense vector (named '')
            # This bypasses LangChain's QdrantVectorStore which can silently fail
            # on hybrid collections (dense '' + sparse 'langchain-sparse').
            results = self.qdrant_client.query_points(
                collection_name=self.collection_name,
                query=query_vector,
                using="",      # explicitly use the unnamed dense vector
                limit=k,
                with_payload=True,
            )
            points = results.points
            logging.debug("[GlobalRecruiterIndex] Qdrant returned %d points.", len(points))

            # Extract unique bot_ids in relevance order
            seen: set = set()
            unique_bot_ids: List[str] = []
            for pt in points:
                payload = pt.payload or {}
                # LangChain stores metadata nested: {"page_content": ..., "metadata": {"bot_id": ...}}
                bot_id = (
                    payload.get("metadata", {}).get("bot_id")
                    or payload.get("bot_id")
                )
                if bot_id and bot_id not in seen:
                    seen.add(bot_id)
                    unique_bot_ids.append(bot_id)

            logging.debug("[GlobalRecruiterIndex] Unique bot_ids found: %s", unique_bot_ids)
            return unique_bot_ids

        except Exception as e:
            logging.exception("[GlobalRecruiterIndex] Qdrant search failed")
            raise RuntimeError("Qdrant search failed.") from e